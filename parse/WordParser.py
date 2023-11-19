from typing import Dict, List

from docx import Document as create_document
from docx.shared import Pt, RGBColor
from docx.document import Document
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH


class WordParser:

    def __init__(self,
                 header: List[str],
                 source: Dict[str, str],
                 title: str = "文档",
                 filter_empty: bool = True) -> None:
        self.doc = create_document()
        self._set_style(self.doc)
        self.title = title
        self.header = header
        self.data = source
        self.filter_empty = filter_empty

    def _set_style(self, doc: Document) -> None:
        doc.styles["Normal"].font.name = u"宋体"
        doc.styles["Normal"]._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        doc.styles["Normal"].font.size = Pt(12)
        doc.styles["Normal"].font.color.rgb = RGBColor(0, 0, 0)
        doc.styles["Heading 1"].font.size = Pt(22)
        doc.styles["Heading 1"].font.color.rgb = RGBColor(0, 0, 0)
        doc.styles["Heading 1"].font.name = u"宋体"
        doc.styles["Heading 1"]._element.rPr.rFonts.set(
            qn('w:eastAsia'), u'宋体')
        doc.styles["Heading 2"].font.size = Pt(16)
        doc.styles["Heading 2"].font.color.rgb = RGBColor(0, 0, 0)
        doc.styles["Heading 2"].font.name = u"宋体"
        doc.styles["Heading 2"]._element.rPr.rFonts.set(
            qn('w:eastAsia'), u'宋体')

    def _add_title(self):
        h1 = self.doc.add_heading(self.title, 1)
        h1.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _add_sub_title(self, title):
        self.doc.add_heading(title, 2)

    def _add_para(self, para):
        self.doc.add_paragraph(para)

    def dump(self, file_path):
        self._add_title()
        for sub_title in self.header:
            if self.filter_empty and not len(self.data[sub_title]):
                continue

            self._add_sub_title(sub_title)
            self._add_para(self.data[sub_title])

        self.doc.save(file_path)
